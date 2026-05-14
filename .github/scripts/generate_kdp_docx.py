#!/usr/bin/env python3
"""python-docx で KDP対応のWordファイルを生成
   pandoc生成版より KDP のアップロード成功率が高い"""
import os
import re
import glob
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

today = datetime.now().strftime('%Y-%m-%d')

# 最新のKindle原稿を取得
manuscript_path = f'商品パッケージ/{today}/kindle_manuscript.md'
if not os.path.exists(manuscript_path):
    candidates = sorted(glob.glob('商品パッケージ/*/kindle_manuscript.md'), reverse=True)
    if candidates:
        manuscript_path = candidates[0]
        print(f"⚠️ 今日の原稿なし。最新を使用: {manuscript_path}")
    else:
        print("❌ 原稿が見つかりません")
        exit(1)

with open(manuscript_path, 'r', encoding='utf-8') as f:
    manuscript = f.read()

print(f"✅ 原稿読み込み: {len(manuscript)}文字")

# 出版ディレクトリ
pub_dir = f'Kindle出版/{today}'
os.makedirs(pub_dir, exist_ok=True)
docx_path = f'{pub_dir}/kdp_manuscript.docx'

# ---- Document 作成 ----
doc = Document()

# 日本語フォント設定
def set_font(run, size=11, bold=False):
    run.font.name = 'ＭＳ 明朝'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'ＭＳ 明朝')
    rFonts.set(qn('w:ascii'), 'ＭＳ 明朝')
    rFonts.set(qn('w:hAnsi'), 'ＭＳ 明朝')
    rPr.append(rFonts)
    run.font.size = Pt(size)
    run.font.bold = bold

# ページ余白
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ---- パース処理 ----
def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=20, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 章の前に改ページ
        p.paragraph_format.page_break_before = True
    elif level == 2:
        set_font(run, size=16, bold=True)
    elif level == 3:
        set_font(run, size=13, bold=True)
    else:
        set_font(run, size=12, bold=True)

def add_paragraph_text(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    # マークダウンの **太字** を処理
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, size=11, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, size=11)

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=11)

# タイトルページ
title_match = re.search(r'^#\s+(.+)', manuscript, re.MULTILINE)
title = title_match.group(1).strip() if title_match else f"AI活用ガイド"

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(200)
title_run = title_p.add_run(title)
set_font(title_run, size=28, bold=True)

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.paragraph_format.space_before = Pt(100)
author_run = author_p.add_run('AIプロンプト実践研究会')
set_font(author_run, size=14)

# タイトルページの後に改ページ
doc.add_page_break()

# ---- 本文パース ----
lines = manuscript.split('\n')
in_code_block = False
skip_first_h1 = True  # 最初のH1はタイトルなので本文ではスキップ

for line in lines:
    stripped = line.strip()

    # コードブロック
    if stripped.startswith('```'):
        in_code_block = not in_code_block
        continue
    if in_code_block:
        continue

    # 空行
    if not stripped:
        continue

    # 見出し
    if stripped.startswith('# '):
        if skip_first_h1:
            skip_first_h1 = False
            continue
        add_heading(stripped[2:], 1)
    elif stripped.startswith('## '):
        add_heading(stripped[3:], 2)
    elif stripped.startswith('### '):
        add_heading(stripped[4:], 3)
    elif stripped.startswith('#### '):
        add_heading(stripped[5:], 4)
    # 箇条書き
    elif stripped.startswith('- ') or stripped.startswith('* '):
        add_bullet(stripped[2:])
    elif re.match(r'^\d+\.\s', stripped):
        add_bullet(re.sub(r'^\d+\.\s', '', stripped))
    # 区切り線
    elif stripped == '---':
        continue
    # 通常段落
    else:
        add_paragraph_text(stripped)

doc.save(docx_path)

# ファイルサイズ確認
size_kb = os.path.getsize(docx_path) / 1024
print(f"✅ KDP用DOCX生成: {docx_path}")
print(f"  サイズ: {size_kb:.1f} KB")
print(f"  タイトル: {title}")
